import os
from openai import OpenAI
from dotenv import load_dotenv
from flask import Flask, Blueprint, make_response, jsonify, request
from flask_restful import Api, Resource
from flask_jwt_extended import jwt_required
from docx import Document

# Create a Flask Blueprint for filesearch
fileSearch_bp = Blueprint('fileSearch_bp', __name__)
api = Api(fileSearch_bp)

load_dotenv()
key = os.getenv("OPENAI_API_KEY")
assistant_id = os.getenv("ASSISTANT_ID")
client = OpenAI(api_key=key)

description = """ 
You are a 'Summarizer Bot' for the 'SME Founders Association (SFA) app'. Your role is to read Zoom transcripts of conversations among SME owners and summarize the content. The goal is to produce a comprehensive report that includes a summary of the conversation and an analysis of the needs expressed by the SME owners. This analysis will help the association tailor events and resources to empower SME owners to run successful businesses.
"""

instructions = """
1. **Read the Transcript**: Carefully read the entire Zoom transcript provided. Ensure you understand the context and main points discussed.

2. **Identify Key Themes**: Identify and list the key themes and topics discussed in the conversation. Pay attention to recurring issues, challenges, and opportunities mentioned by the SME owners.

3. **Summarize the Conversation**: Provide a clear and concise summary of the conversation. Highlight the main points, insights, and any significant remarks made by the participants.

4. **Analyze Needs**: Analyze the needs and concerns expressed by the SME owners. Categorize these needs into different areas such as funding, marketing, mentorship, operations, etc.

5. **Generate Recommendations**: Based on the analysis, generate recommendations for the types of events and resources that the SME Founders Association should consider organizing. These should directly address the identified needs.

6. **Compile the Report**: Compile the summary, needs analysis, and recommendations into a comprehensive report. Ensure the report is well-structured and easy to understand.

7. **Provide Clear Output**: The final output should be a neatly formatted report with sections for:
    - Summary of the Conversation
    - Analysis of Needs
    - Recommendations for Events and Resources

8. **Review and Validate**: Before finalizing the report, review the content to ensure accuracy and completeness. Validate that all significant points have been covered and the recommendations are actionable.

9. **Save the report in .docx formart for a user to be able to download

Remember, your main objective is to provide actionable insights that will help the SME Founders Association in planning effective events and resources for SME owners.
"""

assistant = client.beta.assistants.create(
    name="The SME Insight Bot",
    description=description,
    instructions=instructions,
    model="gpt-4-turbo",
    tools=[{"type": "file_search"}],
)

print(assistant)

vector_store = client.beta.vector_stores.create(name="Transcripts")
print(f"Vector Store Id - {vector_store.id}")

class InteractTranscripts(Resource):
    @jwt_required()
    def post(self):
        data = request.json
        transcripts = data.get('transcripts')

        if not transcripts:
            return make_response(jsonify({"error": "No transcripts provided"}), 400)

        # Prepare the files for upload
        # Prepare the files for upload
        file_streams = []
        try:
            for i, transcript in enumerate(transcripts):
                file_path = f"/tmp/transcript_{i}.txt"

                # Check if transcript is a dictionary, extract text if necessary
                if isinstance(transcript, dict):
                    transcript_text = transcript.get("content")  # Assuming "text" is the key
                    if not transcript_text:
                        return make_response(jsonify({"error": "Transcript missing content"}), 400)
                else:
                    transcript_text = transcript

                # Write the transcript to a temporary file
                with open(file_path, "w") as f:
                    f.write(transcript_text)
                
                # Open the file in binary read mode and append to the file_streams list
                file_streams.append(open(file_path, "rb"))
            
            # Add Files to Vector Store
            file_batch = client.beta.vector_stores.file_batches.upload_and_poll(
                vector_store_id=vector_store.id,
                files=file_streams
            )

            # Print the status and the file counts of the batch to see the result of this operation
            print(file_batch.status)
            print(file_batch.file_counts)
        finally:
            # Clean up: Close file streams and remove temporary files
            for file_stream in file_streams:
                file_stream.close()
                
            for i, _ in enumerate(transcripts):
                file_path = f"/tmp/transcript_{i}.txt"
                if os.path.exists(file_path):
                    os.remove(file_path)

        
        #update the Assitant with a vectore store
        assistant=client.beta.assistants.update(
            assistant_id=assistant_id,
            tool_resources={"file_search":{"vector_store_ids":[vector_store.id]}}
        )
        print("Assistant Updated with Vector Store")

        #Create a Thread
        thread=client.beta.threads.create()
        print(f"Your thread is {thread.id}\n\n")

        #run a loop where user can ask questions
        while True:
            text=input("What's your question?\n")
            message=client.beta.threads.messages.create(
                thread_id=thread.id,
                role="user",
                content=text,
            )
            run=client.beta.threads.runs.create_and_poll(
                thread_id=thread.id,assistant_id=assistant.id
            )
            messages=list(client.beta.threads.messages.list(thread_id=thread.id,run_id=run.id))
            message_content=messages[0].content[0].text
            print("Response:\n")
            print(f"{message_content.value}\n")            

       

           

# Register the ProcessTranscripts resource with the API
api.add_resource(InteractTranscripts, '/interact_transcripts')

class ProcessTranscripts(Resource):
    @jwt_required()
    def post(self):
        data = request.json
        transcription = data.get('transcripts')

        if not transcription:
            return make_response(jsonify({"error": "No transcripts provided"}), 400)

        try:
            # Process the transcription to extract meeting minutes
            minutes = self.meeting_minutes(transcription)

            # Save the minutes to a DOCX file
            filename = 'meeting_minutes.docx'
            self.save_as_docx(minutes, filename)

            # Return a success message along with the filename
            return make_response(jsonify({"message": "Transcription processed successfully", "filename": filename}), 200)

        except Exception as e:
            return make_response(jsonify({"error": str(e)}), 500)

    def abstract_summary_extraction(self, transcription):
        response = client.chat.completions.create(
            model="gpt-4-turbo",
            temperature=0,
            messages=[
                {
                    "role": "system",
                    "content": "You are a highly skilled AI trained in language comprehension and summarization. I would like you to read the following text and summarize it into a concise abstract paragraph. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points."
                },
                {
                    "role": "user",
                    "content": transcription
                }
            ]
        )
        return response.choices[0].message.content

    def key_points_extraction(self, transcription):
        response = client.chat.completions.create(
            model="gpt-4-turbo",
            temperature=0,
            messages=[
                {
                    "role": "system",
                    "content": "You are a proficient AI with a specialty in distilling information into key points. Based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about."
                },
                {
                    "role": "user",
                    "content": transcription
                }
            ]
        )
        return response.choices[0].message.content

    def action_item_extraction(self, transcription):
        response = client.chat.completions.create(
            model="gpt-4-turbo",
            temperature=0,
            messages=[
                {
                    "role": "system",
                    "content": "You are an AI expert in analyzing conversations and extracting action items. Please review the text and identify any tasks, assignments, or actions that were agreed upon or mentioned as needing to be done. These could be tasks assigned to specific individuals, or general actions that the group has decided to take. Please list these action items clearly and concisely."
                },
                {
                    "role": "user",
                    "content": transcription
                }
            ]
        )
        return response.choices[0].message.content
    
    def sentiment_analysis(self, transcription):
        response = client.chat.completions.create(
            model="gpt-4-turbo",
            temperature=0,
            messages=[
                {
                    "role": "system",
                    "content": "As an AI with expertise in language and emotion analysis, your task is to analyze the sentiment of the following text. Please consider the overall tone of the discussion, the emotion conveyed by the language used, and the context in which words and phrases are used. Indicate whether the sentiment is generally positive, negative, or neutral, and provide brief explanations for your analysis where possible."
                },
                {
                    "role": "user",
                    "content": transcription
                }
            ]
        )
        return response.choices[0].message.content

    def meeting_minutes(self, transcription):
        abstract_summary = self.abstract_summary_extraction(transcription)
        key_points = self.key_points_extraction(transcription)
        action_items = self.action_item_extraction(transcription)
        sentiment = self.sentiment_analysis(transcription)
        return {
            'abstract_summary': abstract_summary,
            'key_points': key_points,
            'action_items': action_items,
            'sentiment': sentiment
        }

    def save_as_docx(self, minutes, filename):
        doc = Document()
        for key, value in minutes.items():
            # Replace underscores with spaces and capitalize each word for the heading
            heading = ' '.join(word.capitalize() for word in key.split('_'))
            doc.add_heading(heading, level=1)
            doc.add_paragraph(value)
            # Add a line break between sections
            doc.add_paragraph()
        doc.save(filename)

api.add_resource(ProcessTranscripts, '/process_transcripts')